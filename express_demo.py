# 0. 删除原output文件夹,创建新output文件夹
# 1. 复制 public\template的文件到output文件夹下
# 2. 遍历文件夹下所有docx文件，替换关键字(关键字从config.json中读取)
# 3. 打包 output文件夹

import os
import shutil
from docx import Document
import json

# 0.删除output 文件夹
print("**********************0. 删除原output文件夹,创建新output文件夹************************")
filepath = r"D:\project\pythonWps\public\output"
if os.path.exists(filepath):
    shutil.rmtree(filepath,True)
    print("dir "+ str(filepath)+ " removed!")
else:
    print("dir "+ str(filepath) + " not exist.")

# 0.创建output 文件夹
os.mkdir(filepath)
print("dir "+ str(filepath) + " create.")

# 1.复制 public\template的文件到output文件夹下
print("**********************1.复制 public\template的文件到output文件夹下************************")
src = r"D:\project\pythonWps\public\template"
dst = r"D:\project\pythonWps\public\output"

def get_filelist(dir):
    Filelist = []
    DirList = []
    for home, dirs, files in os.walk(dir):
        for filename in files:
            # 文件名列表，包含完整路径
            if os.path.isfile(os.path.join(home, filename)):
                DirList.append(os.path.join(home, filename))
            # 文件名列表，只包含文件名
            Filelist.append( filename)
    return Filelist,DirList

Result = get_filelist(src)
Files = Result[0]
Dirs = Result[1]
for (file,Dir) in zip(Files,Dirs):
    print("copy " + str(os.path.join(src,file)) + "to " + str(os.path.join(dst,file)))
    shutil.copyfile(os.path.join(src,file),os.path.join(dst,file))

# 2.遍历文件夹下所有docx文件，替换关键字(关键字从config.json中读取)
print("**********************2.遍历文件夹下所有docx文件，替换关键字(关键字从config.json中读取)************************")
OutputFiles,OutputFileDirs = get_filelist(dst)
'''
全局内容替换
请确保要替换的内容样式一致

Args:
    doc: 文档对象
    old_text: 要被替换的文本
    new_text: 要替换成的文本
'''
def replace_text(doc, old_text, new_text):
    # 遍历每个段落
    for p in doc.paragraphs:
        # 如果要搜索的内容在该段落
        if old_text in p.text:
            # 使用 runs 替换内容但不改变样式
            # 注意！runs 会根据样式分隔内容，确保被替换内容的样式一致
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
    # 遍历每个表格               
    for table in document.tables:
        #遍历行
        for i in range(0,len(table.rows)):
        #遍历列，如合并单元格，这里会出现重复
            for j in range(0,len(table.columns)):
                if old_text in table.cell(i,j).text:
                    table.cell(i,j).paragraphs[0].runs[0].text = new_text
'''
从json中读取关键字的值

Args:
    key: 关键字
'''
def get_json_keyValue(key):
    # 获取json里面数据
 
    with open('config.json', 'rb') as f:
        # 定义为只读模型，并定义名称为f
 
        params = json.load(f)
        # 加载json文件中的内容给params

        keyValue = params['data'][0][key]
    f.close()
    # 关闭json读模式
    return keyValue

project_name_old = get_json_keyValue('project_name_old')
project_code_old = get_json_keyValue('project_code_old')
project_date_old = get_json_keyValue('project_date_old')
project_position1_old = get_json_keyValue('project_position1_old')
project_position2_old = get_json_keyValue('project_position2_old')
project_object_old = get_json_keyValue('project_object_old')
project_time_old = get_json_keyValue('project_time_old')
name_old = get_json_keyValue('name_old')
project_date_month_old = get_json_keyValue('project_date_month_old')

project_name_new = get_json_keyValue('project_name_new')
project_code_new = get_json_keyValue('project_code_new')
project_date_new = get_json_keyValue('project_date_new')
project_position1_new = get_json_keyValue('project_position1_new')
project_position2_new = get_json_keyValue('project_position2_new')
project_object_new = get_json_keyValue('project_object_new')
project_time_new = get_json_keyValue('project_time_new')
name_new = get_json_keyValue('name_new')
project_date_month_new = get_json_keyValue('project_date_month_new')


print("将 "+str(project_name_old)+" 替换为 "+str(project_name_new))
print("将 "+str(project_code_old)+" 替换为 "+str(project_code_new))
print("将 "+str(project_date_old)+" 替换为 "+str(project_date_new))
print("将 "+str(project_position1_old)+" 替换为 "+str(project_position1_new))
print("将 "+str(project_position2_old)+" 替换为 "+str(project_position2_new))
print("将 "+str(project_object_old)+" 替换为 "+str(project_object_new))
for OutputFile in OutputFileDirs:
    print(OutputFile)
    # 从文件创建文档对象
    document = Document(OutputFile);
    replace_text(document,project_name_old,project_name_new)
    replace_text(document,project_code_old,project_code_new)
    replace_text(document,project_date_old,project_date_new)
    replace_text(document,project_position1_old,project_position1_new)
    replace_text(document,project_position2_old,project_position2_new)
    replace_text(document,project_object_old,project_object_new)
    replace_text(document,project_time_old,project_time_new)
    replace_text(document,name_old,name_new)
    replace_text(document,project_date_month_old,project_date_month_new)
    # 保存文档
    FileName = OutputFile.split(os.path.sep)[5]
    print("FileName:",FileName)
    document.save(OutputFile);
# 3. 打包 output文件夹

# 将path_1处的文件归档到path_2处
path_1 = r"D:\project\pythonWps\public\output"
path_2 = r"D:\project\pythonWps\public\download"
new_path = shutil.make_archive(path_2, 'zip', path_1)
print(new_path)