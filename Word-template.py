# 导入
from docx import Document

'''
全局内容替换
请确保要替换的内容样式一致

Args:
    doc: 文档对象
    old_text: 要被替换的文本
    new_text: 要替换成的文本
'''
def replace_text(doc, old_text, new_text):
    # 遍历每个段落
    for p in doc.paragraphs:
        # 如果要搜索的内容在该段落
        if old_text in p.text:
            # 使用 runs 替换内容但不改变样式
            # 注意！runs 会根据样式分隔内容，确保被替换内容的样式一致
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
    # 遍历每个表格               
    for table in document.tables:
        #遍历行
        for i in range(0,len(table.rows)):
        #遍历列，如合并单元格，这里会出现重复
            for j in range(0,len(table.columns)):
                if old_text in table.cell(i,j).text:
                    table.cell(i,j).paragraphs[0].runs[0].text = new_text


# 从文件创建文档对象
document = Document('./表格/0.摘要.docx');
#查找替换
replace_text(document, '[FJNPLYZZZB2020022]', 'FJNPLYZZZB2020023')
# 保存文档
document.save('demo.docx');